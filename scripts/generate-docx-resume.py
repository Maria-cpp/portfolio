"""
Generate Maria Naseem's Senior AI Engineer resume as a .docx file.
Usage: python scripts/generate-docx-resume.py
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_cell_border(cell, **kwargs):
    """Set cell border. Pass side='val,sz,color' e.g. top={'val':'single','sz':'4','color':'999999'}"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
    for edge, attrs in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{attrs["val"]}" w:sz="{attrs["sz"]}" '
            f'w:space="0" w:color="{attrs["color"]}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def remove_table_borders(table):
    """Remove all borders from a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def add_horizontal_rule(doc):
    """Add a thin horizontal rule via a paragraph with bottom border."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="999999"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def add_section_heading(doc, text):
    """Add an uppercase bold section heading."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor(0, 0, 0)


def set_run_font(run, size=Pt(9.5), bold=False, italic=False, name="Arial", color=None):
    run.font.name = name
    run.font.size = size
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def body_paragraph(doc, text="", size=Pt(9.5), bold=False, alignment=None, space_after=Pt(2)):
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = space_after
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold)
    return p


def add_bullet(doc, text, tech_suffix=None):
    """Add a bullet point. If tech_suffix provided, append it in italic."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)

    if tech_suffix:
        main_text = text
        run = p.add_run(main_text)
        set_run_font(run, size=Pt(9))
        tech_run = p.add_run(f" Tech: {tech_suffix}")
        set_run_font(tech_run, size=Pt(9), italic=True)
    else:
        run = p.add_run(text)
        set_run_font(run, size=Pt(9))
    # Clear default run if any
    if p.runs and p.runs[0].text == "" and len(p.runs) > 1:
        pass
    return p


def add_job_header(doc, title, period):
    """Add job title (left) and period (right) using a borderless table."""
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(table)

    # Set full width
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:type="pct" w:w="5000"/>')
    tblPr.append(tblW)

    left_cell = table.cell(0, 0)
    right_cell = table.cell(0, 1)

    # Title
    lp = left_cell.paragraphs[0]
    lp.paragraph_format.space_before = Pt(4)
    lp.paragraph_format.space_after = Pt(0)
    run = lp.add_run(title)
    set_run_font(run, size=Pt(10), bold=True)

    # Period
    rp = right_cell.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.paragraph_format.space_before = Pt(4)
    rp.paragraph_format.space_after = Pt(0)
    run = rp.add_run(period)
    set_run_font(run, size=Pt(10))


def add_company_line(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    set_run_font(run, size=Pt(9))


# ---------------------------------------------------------------------------
# Main document builder
# ---------------------------------------------------------------------------

def build_resume():
    doc = Document()

    # --- Page setup ---
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # --- Reduce default paragraph spacing ---
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(9.5)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(2)

    # Fix List Bullet style
    if "List Bullet" in [s.name for s in doc.styles]:
        lb = doc.styles["List Bullet"]
        lb.font.name = "Arial"
        lb.font.size = Pt(9)

    # ===== NAME =====
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("MARIA NASEEM")
    set_run_font(run, size=Pt(20), bold=True)

    # ===== HEADLINE =====
    p = body_paragraph(doc, "Senior AI Engineer | AI Solutions Architect",
                       size=Pt(13), bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(2))

    # ===== TAGLINE =====
    body_paragraph(doc, "Agentic AI \u00b7 LLMs \u00b7 Computer Vision | Enterprise AI",
                   size=Pt(10), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(2))

    # ===== CONTACT =====
    body_paragraph(
        doc,
        "Islamabad, Pakistan | marianaseem99@gmail.com | LinkedIn: linkedin.com/in/maria-naseem/ "
        "| GitHub: github.com/Maria-cpp | Portfolio: maria-naseem.vercel.app",
        size=Pt(9), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(2)
    )

    add_horizontal_rule(doc)

    # ===== HIGHLIGHTS =====
    highlights = [
        ("\u2714 10+ Years Professional Experience", "\u2714 Computer Vision"),
        ("\u2714 4+ Years AI Engineering", "\u2714 Agentic AI & Multi-Agent Systems"),
        ("\u2714 12+ Production Systems", "\u2714 LLM Applications & Real-Time AI"),
        ("\u2714 Enterprise & Government AI Solutions", "\u2714 FastAPI Expert \u00b7 Azure Deployment"),
    ]
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(table)
    tbl = table._tbl
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:type="pct" w:w="5000"/>')
    tbl.tblPr.append(tblW)

    for i, (left, right) in enumerate(highlights):
        for j, text in enumerate((left, right)):
            cell = table.cell(i, j)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(text)
            set_run_font(run, size=Pt(9.5))

    add_horizontal_rule(doc)

    # ===== PROFESSIONAL SUMMARY =====
    add_section_heading(doc, "Professional Summary")
    body_paragraph(
        doc,
        "Senior AI Engineer who designs and ships production AI systems for enterprise and "
        "government clients. Architected real-time video analytics deployed on Azure, a four-stage "
        "facial recognition attendance system running live on-premises, and a bank-grade data "
        "protection vault \u2014 each delivered from requirements through on-site client demonstrations. "
        "Built 12+ AI systems spanning computer vision, agentic workflows, and LLM applications. "
        "Currently leading AI platform engineering at Arwen Tech and running Forward Deployed "
        "Engineering engagements through ZumfluxAI.",
        size=Pt(9.5), space_after=Pt(4)
    )

    add_horizontal_rule(doc)

    # ===== CORE COMPETENCIES =====
    add_section_heading(doc, "Core Competencies")
    competencies = [
        ("Agentic AI", "LLM Applications", "Multi-Agent Systems"),
        ("Computer Vision", "RAG Pipelines", "Real-Time AI"),
        ("End-to-End Architecture", "FastAPI", "Azure Deployment"),
        ("Docker & CI/CD", "Python", "PostgreSQL"),
    ]
    table = doc.add_table(rows=4, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(table)
    tbl = table._tbl
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:type="pct" w:w="5000"/>')
    tbl.tblPr.append(tblW)

    for i, row_data in enumerate(competencies):
        for j, text in enumerate(row_data):
            cell = table.cell(i, j)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(text)
            set_run_font(run, size=Pt(9.5))

    add_horizontal_rule(doc)

    # ===== TECHNICAL SKILLS =====
    add_section_heading(doc, "Technical Skills")
    skills = [
        ("AI & ML:", " Agentic AI, LLM Applications, Multi-Agent Systems, RAG, Computer Vision, NLP, OCR"),
        ("Frameworks:", " LangGraph, LangChain, OpenAI SDK, Anthropic SDK, FastAPI, YOLO/Ultralytics, OpenCV"),
        ("Backend:", " Python, PostgreSQL, Redis, Celery, WebSockets"),
        ("Cloud & DevOps:", " Docker, Azure, GitHub Actions, Linux, Nginx"),
        ("Frontend:", " Next.js, React, TypeScript, Tailwind CSS"),
    ]
    for label, items in skills:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(label)
        set_run_font(run, size=Pt(9), bold=True)
        run = p.add_run(items)
        set_run_font(run, size=Pt(9))

    add_horizontal_rule(doc)

    # ===== PROFESSIONAL EXPERIENCE =====
    add_section_heading(doc, "Professional Experience")

    # --- Job 1 ---
    add_job_header(doc, "AI Solutions Architect & Full-Stack Developer", "2025 \u2014 Present")
    add_company_line(doc, "Arwen Tech \u00b7 Islamabad, Pakistan")

    add_bullet(
        doc,
        "Designed end-to-end architecture for real-time production counting using YOLO and OpenVINO. "
        "Built RTSP pipeline with object tracking, tamper detection, and FastAPI microservices. "
        "Containerized deployment on Azure. Successfully demonstrated solution at client production facility."
    )
    add_bullet(
        doc,
        "Engineered a four-stage facial recognition attendance system processing multi-camera live "
        "video streams with human verification gates. Architecture scales to thousands of employees. "
        "Deployed live on-premises for real-world training and validation."
    )
    add_bullet(
        doc,
        "Built enterprise Security Vault protecting sensitive customer information through tokenization "
        "and encryption with immutable audit logging. 86 automated tests.",
        tech_suffix="FastAPI, PostgreSQL, Redis."
    )
    add_bullet(
        doc,
        "Shipped an MCP-native observability platform with AI-powered alert analysis performing "
        "root-cause inference and severity classification. Built and deployed the MCP server for tool integration.",
        tech_suffix="Prometheus, Grafana, Claude SDK."
    )
    add_bullet(
        doc,
        "Delivered a multilingual NLP platform analyzing public feedback across 8+ low-resource "
        "languages with automated sentiment analysis and AI briefing generation. 107 tests."
    )

    # --- Job 2 ---
    add_job_header(doc, "Corporate Affairs, Legal Operations & AI Solutions Engineer", "2024 \u2014 2025")
    add_company_line(doc, "Green Tourism Pvt. Ltd. \u00b7 Islamabad, Pakistan")

    add_bullet(
        doc,
        "Identified operational bottleneck and built an AI contract intelligence platform automating "
        "legal operations \u2014 contract ingestion, semantic search, metadata extraction, and workflow "
        "orchestration for a national tourism company.",
        tech_suffix="LangGraph, Claude, RAG, pgvector."
    )
    add_bullet(
        doc,
        "Automated meeting scheduling, minutes management, follow-up generation, and contract "
        "lifecycle workflows with audit logging."
    )
    add_bullet(
        doc,
        "Served as primary legal and regulatory liaison \u2014 managing compliance, Board of Directors "
        "coordination, and cross-functional stakeholder communication."
    )
    add_bullet(
        doc,
        "Delivered corporate website and e-commerce platform with custom development and API integrations."
    )

    # ===== PAGE BREAK =====
    doc.add_page_break()

    # ===== PROFESSIONAL EXPERIENCE (continued) =====
    add_section_heading(doc, "Professional Experience (continued)")

    # --- Job 3 ---
    add_job_header(doc, "Co-Founder \u00b7 Full-Stack Developer", "2023 \u2014 2025")
    add_company_line(doc, "ZUM Services Providers \u00b7 Pakistan")

    add_bullet(
        doc,
        "Co-founded and built the full-stack platform end-to-end \u2014 backend, database, APIs, and "
        "frontend \u2014 while running business operations, client acquisition, and service delivery."
    )
    add_bullet(
        doc,
        "Containerized all services with Docker and CI/CD via GitHub Actions for production deployments."
    )

    # --- Job 4 ---
    add_job_header(doc, "Blockchain Developer", "2020 \u2014 2022")
    add_company_line(doc, "MediaPark \u00b7 Pakistan")

    add_bullet(
        doc,
        "Designed native blockchain systems and decentralized applications with cryptographic "
        "signing and peer-to-peer networking."
    )
    add_bullet(
        doc,
        "Built compliance-oriented smart contract frameworks with access controls and audit documentation."
    )

    # --- Job 5 ---
    add_job_header(doc, "Previous Experience", "2016 \u2014 2020")
    add_company_line(doc, "IBM Pakistan (Intern) \u00b7 Prime Tele Power \u00b7 Team Work Construction")
    add_bullet(doc, "Software Development \u00b7 Operations \u00b7 IT Support")

    add_horizontal_rule(doc)

    # ===== SELECTED PROJECTS =====
    add_section_heading(doc, "Selected Projects")

    projects = [
        (
            "AI Video Analytics Platform",
            " \u2014 Designed real-time production counting architecture for regulated-industry client. "
            "Deployed on Azure, demonstrated on-site.",
            "YOLO, OpenCV, FastAPI, Azure."
        ),
        (
            "Facial Recognition Attendance",
            " \u2014 Four-stage pipeline with human verification, multi-camera RTSP, deployed on-premises.",
            "DeepFace, OpenCV, FAISS."
        ),
        (
            "Security Vault",
            " \u2014 Enterprise data protection through tokenization and encryption with immutable audit trails.",
            "FastAPI, PostgreSQL, Redis."
        ),
        (
            "Multilingual NLP Platform",
            " \u2014 Analyzed feedback across 8+ low-resource languages with AI briefing generation.",
            "FastAPI, Kafka, PyTorch."
        ),
        (
            "Agentic Contract Intelligence",
            " \u2014 Multi-agent system automating contract lifecycle for a national tourism company.",
            "LangGraph, Claude, RAG."
        ),
    ]

    for title, desc, tech in projects:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(title)
        set_run_font(run, size=Pt(9.5), bold=True)
        run = p.add_run(desc)
        set_run_font(run, size=Pt(9.5))
        run = p.add_run(f" Tech: {tech}")
        set_run_font(run, size=Pt(9.5), italic=True)

    add_horizontal_rule(doc)

    # ===== EDUCATION =====
    add_section_heading(doc, "Education")
    body_paragraph(doc, "Master of Information Technology \u2014 Quaid-i-Azam University", size=Pt(9.5))

    add_horizontal_rule(doc)

    # ===== CERTIFICATIONS =====
    add_section_heading(doc, "Certifications")
    body_paragraph(
        doc,
        "Agentic AI Level 2 \u00b7 Agentic AI Level 1 \u00b7 IBM Build RAG Applications \u00b7 "
        "DeepLearning AI \u00b7 Coursera GenAI",
        size=Pt(9.5)
    )

    return doc


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    # Ensure output directory exists
    output_dir = os.path.join("public", "resumes")
    os.makedirs(output_dir, exist_ok=True)

    output_path = os.path.join(output_dir, "Maria_Naseem_Senior_AI_Engineer.docx")

    doc = build_resume()
    doc.save(output_path)
    print(f"Resume generated successfully: {output_path}")
